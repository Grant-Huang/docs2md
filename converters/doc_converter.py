"""
Word 文档转 Markdown 转换器
支持 .docx 格式，.doc 旧格式暂不支持
"""
import sys
import os
import io
import re
import json

# 设置标准输出和标准错误的编码为utf-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')


def _send_json(obj):
    """发送 JSON 格式的 SSE 消息"""
    print(f'data: {json.dumps(obj, ensure_ascii=False)}\n\n', flush=True)


def _dedupe_row(cells):
    """合并连续重复单元格（处理 Word 合并单元格导致的重复列）"""
    if not cells:
        return cells
    result = [cells[0]]
    for c in cells[1:]:
        if c != result[-1]:
            result.append(c)
    return result


def _table_to_markdown(table):
    """将 docx 表格转换为 markdown 表格"""
    rows = []
    for row in table.rows:
        cells = [
            cell.text.strip().replace('\n', ' ').replace('|', '\\|')
            for cell in row.cells
        ]
        cells = _dedupe_row(cells)
        if cells:
            rows.append('| ' + ' | '.join(cells) + ' |')
    if not rows:
        return ''
    col_count = max(len(r.split('|')) - 2 for r in rows) or 1
    sep = '| ' + ' | '.join(['---'] * col_count) + ' |'
    return rows[0] + '\n' + sep + '\n' + '\n'.join(rows[1:])


def convert_doc_to_text(doc_path, output_format='markdown'):
    """将 Word 文档转换为文本，支持 markdown 和纯文本格式"""
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')

        if not os.path.exists(doc_path):
            _send_json({"type": "error", "content": f"文件不存在: {doc_path}"})
            return None

        _, ext = os.path.splitext(doc_path)
        if ext.lower() == '.doc':
            _send_json({
                "type": "error",
                "content": "不支持 .doc 格式，请使用 .docx 格式或先将文件另存为 .docx"
            })
            return None

        file_size = os.path.getsize(doc_path)
        _send_json({"type": "debug", "content": f"文件大小: {file_size} 字节"})

        if file_size == 0:
            _send_json({"type": "error", "content": "文件为空"})
            return None

        try:
            from docx import Document
        except ImportError:
            _send_json({
                "type": "error",
                "content": "缺少 python-docx 库，请运行: pip install python-docx"
            })
            return None

        _send_json({"type": "debug", "content": "正在打开 Word 文档..."})

        doc = Document(doc_path)
        output_parts = []

        def process_blocks(blocks):
            """处理段落和表格块"""
            for block in blocks:
                if hasattr(block, 'text'):  # Paragraph
                    para = block
                    text = para.text.strip()
                    if not text:
                        if output_format == 'markdown':
                            output_parts.append('')
                        continue
                    if output_format == 'text':
                        output_parts.append(text)
                    else:
                        style_name = para.style.name if para.style else ''
                        if style_name.startswith('Heading'):
                            try:
                                match = re.search(r'\d+', style_name)
                                level = int(match.group()) if match else 1
                                level = min(max(level, 1), 6)
                                output_parts.append(f"{'#' * level} {text}")
                            except (ValueError, AttributeError):
                                output_parts.append(text)
                        else:
                            output_parts.append(text)
                else:  # Table
                    table = block
                    if output_format == 'markdown':
                        md_table = _table_to_markdown(table)
                        if md_table:
                            output_parts.append(md_table)
                            output_parts.append('')
                    else:
                        for row in table.rows:
                            cells = [
                                c.text.strip().replace('\n', ' ')
                                for c in row.cells
                            ]
                            output_parts.append('\t'.join(cells))

        try:
            # 按文档顺序遍历段落和表格
            from docx.document import Document as _Document
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.table import CT_Tbl
            from docx.table import Table
            from docx.text.paragraph import Paragraph

            def iter_block_items(parent):
                if isinstance(parent, _Document):
                    parent_elm = parent.element.body
                else:
                    parent_elm = parent._element
                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            blocks = list(iter_block_items(doc))
        except Exception as e:
            _send_json({
                "type": "debug",
                "content": "按顺序解析失败，使用备用方式: " + str(e)
            })
            blocks = list(doc.paragraphs) + list(doc.tables)

        process_blocks(blocks)

        # 提取页眉页脚
        header_footer_parts = []
        for section in doc.sections:
            for attr in ('header', 'first_page_header', 'even_page_header',
                         'footer', 'first_page_footer', 'even_page_footer'):
                try:
                    hf = getattr(section, attr, None)
                    if hf is not None:
                        for para in hf.paragraphs:
                            t = para.text.strip()
                            if t:
                                header_footer_parts.append(t)
                except (AttributeError, KeyError):
                    pass
        if header_footer_parts:
            header_text = '\n'.join(dict.fromkeys(header_footer_parts))
            if output_format == 'markdown':
                output_parts.insert(0, '## 页眉页脚\n\n' + header_text)
            else:
                output_parts.insert(0, '=== 页眉页脚 ===\n' + header_text)

        final_text = '\n\n'.join(
            p for p in output_parts if p is not None
        ).strip()
        final_text = re.sub(r'\n\s*\n\s*\n', '\n\n', final_text)

        _send_json({
            "type": "debug",
            "content": f"提取文本长度: {len(final_text)} 字符"
        })

        if final_text:
            _send_json({"type": "complete", "content": final_text})
            return final_text
        else:
            _send_json({
                "type": "error",
                "content": "警告：没有提取到任何文本"
            })
            return None

    except Exception as e:
        error_msg = f"转换过程中出错: {str(e)}"
        _send_json({"type": "error", "content": error_msg})
        import traceback
        _send_json({
            "type": "error",
            "content": f"错误详情: {traceback.format_exc()}"
        })
        return None


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(
            "Usage: python doc_converter.py <doc_file_path> [format]",
            file=sys.stderr
        )
        print("format: 'markdown' (default) or 'text'", file=sys.stderr)
        sys.exit(1)

    doc_path = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else 'markdown'

    if output_format not in ['markdown', 'text']:
        print(
            "Error: format must be either 'markdown' or 'text'",
            file=sys.stderr
        )
        sys.exit(1)

    result = convert_doc_to_text(doc_path, output_format)
    if result:
        sys.exit(0)
    else:
        sys.exit(1)
